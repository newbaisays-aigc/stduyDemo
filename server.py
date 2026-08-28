#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
server.py — 统一 Web 服务（端口 8000）
同时提供：
  - 静态文件托管（V2.2 开源演示版页面及配套 js/图片等，等价于 python -m http.server）
  - POST /api/clean  题库「清洗矫正」的大模型清洗（MiniMax 中国版）
  - GET  /health
前端使用 location.origin + "/api/clean"（同源同端口），彻底避免跨端口防火墙/回环问题。
"""
import argparse
import json
import os
import re
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

try:
    import yaml
except Exception:
    yaml = None
try:
    import httpx
except Exception:
    httpx = None

ROOT = Path(__file__).resolve().parent
INDEX = "people-knowledge-graph-V2.2-open.html"
DEFAULT_CONFIG = Path(os.environ.get("MAVIS_CONFIG", str(Path.home() / ".mavis" / "config.yaml")))
MINIMAX_CHINA_BASE = "https://api.minimaxi.com"
DEFAULT_MODEL_ID = "MiniMax-M3"

SYSTEM_PROMPT = (
    "你是一名资深语文教研助手。请对用户粘贴的试卷或书本内容做清洗与矫正，规则：\n"
    "1) 修正换行：把并列条目（数字序号、①圈码、字母选项 A./B. 等）各自独立成行；把句子中间被拆断的词语合并回。\n"
    "2) 修正标点：汉字后的半角标点统一为全角（, . ? ! ; : → ，。？！；：），去掉多余空格与连续空行。\n"
    "3) 保留结构：内容标题（如「重要词语解释」「多音字注音并组词」）独立成行；词条、解析等作为其下的正文。\n"
    "4) 不要修改事实与文字本身的含义，只做排版与格式矫正。\n"
    "5) 直接输出矫正后的纯文本，不要解释、不要加引号包裹、不要输出额外说明。"
)
STRUCT_PROMPT = (
    "你是一名资深语文教研助手。用户会给一段试卷题目原文；原文里的加点/波浪线/下划线等标记已经用占位符表示（形如 ◉dot1◉、◉wave2◉、◉u3◉），这些占位符属于题目本身，必须一个都不落地原样保留在输出的对应位置。\n"
    "请完整重建这段题目清单：\n"
    "1) 逐题原样保留题目文字——序号（如 1．/（1）/2．）、填空下划线 ________、分值（如（18分））、标点、书名号《》、标题等一律保留；不要删改、不要重写题目措辞、不要改任何一个 ◉...◉ 占位符。\n"
    "2) 在每一道题的原文正下方，先写一行「答案：」并把该题答案写在「答案：」之后；若该题还带解析/讲解/原文出处，再单独一行「解析：」+解析。\n"
    "3) 讲次/大题标题、配分说明作为标题行原样保留，但不要给它们加「答案：」/「解析：」。\n"
    "4) 不要修改题目、答案与解析的事实内容；输出前逐字核对：每个 ◉...◉ 占位符都必须完整出现在对应题目位置，绝不丢失、绝不改写成普通文字。\n"
    "5) 直接输出“题目原文（含全部占位符）+ 每题下方 答案：/解析：”的完整文本，保持原有换行与排版；不要解释，不要加引号包裹，不要输出额外说明。"
)
# ANSWERS_PROMPT：仅补充单题答案与解析，绝不重建/改写题目原文。
# 前端已逐题把「题目原文（含占位符）」拆成小块调用；题目原文由前端本地 100% 保留，
# 模型只负责针对该题给出答案与解析，输出不改动题目文字。
ANSWERS_PROMPT = (
    "你是一名资深语文教研助手。用户会给一道语文题目（可能包含多个小题），题目文字里可能带有加点/波浪线/下划线占位符（形如 ◉dot1◉、◉wave2◉、◉u3◉），它们属于题目的一部分。\n"
    "请只针对这道题给出答案与解析：\n"
    "1) 第一行写「答案：」，随后写出本题（含各小题）的准确答案；若题目本身已是完整陈述、没有需要作答的填空，则写「答案：见题干，无需作答」或类似说明即可，不要凭空编造。\n"
    "2) 若题目还带解析/讲解/出处，再单独一行写「解析：」+简要解析；没有把握或没有解析时可不写「解析：」一行，切勿编造。\n"
    "3) 绝对不要重复、改写、重排任何题目原文；只输出「答案：」和「解析：」两行（解析行可选）。输出中如确实需要引用题目里的占位符字样，原样保留 ◉...◉，但通常答案文字里不需要包含它们。\n"
    "4) 直接输出纯文本，不要解释，不要加引号包裹，不要输出任何额外说明。"
)
# FORMAT_PROMPT：把一段试卷/题目内容规范成统一排版格式(章节/标题/正文缩进/选项/换行)。
# 只做格式规范，绝不改动题目、答案、解析的事实内容与占位符。
FORMAT_PROMPT = (
    "你是一名资深语文教研助手，负责把用户粘贴的试卷/题目内容排成统一、规范的格式。请只做排版规范，不要改任何事实内容。\n"
    "统一格式规则：\n"
    "1) 章节标题：以「# 」开头（如“# 文言文专项”）；一级标题「## 」；二级标题「### 」；三级标题「#### 」。\n"
    "2) 题目编号：用「1．」「2．」等，独立成行；子问用「（1）（2）」按行缩进。\n"
    "3) 选项：用「A．B．C．D．」每个选项独立成行，紧跟在题干后。\n"
    "4) 正文：普通段落直接写，不需要前缀；原文引文、较长的阅读材料可用行首「> 」表示缩进段。\n"
    "5) 换行：每一条(标题/题号/选项/答案/解析)各占一行，不要挤在一行；空行适当分隔。\n"
    "6) 「答案：」「解析：」各自独立成行，放在对应题目下方，内容照抄不改。\n"
    "7) 保留所有加点/波浪线/下划线占位符：原样输出 ◉dot1◉、◉wave2◉、◉u3◉ 等，绝不丢失、绝不改写成普通文字；也保留 [[dot:…]]、**…** 这类原生标记原样输出（不要替换成其它符号）。\n"
    "8) 逐字核对事实内容：题目、答案、解析、分值（如（3分））、书名号《》、加点词一个都不要改动或丢失。\n"
    "直接输出规范排版后的纯文本，不要解释、不要加引号包裹、不要输出额外说明。"
)
# EXPLAIN_PROMPT：题目答案已在源文档给出，只需针对该题写出「解析」——
# 解析内容是答题采分点的解释、如何拿分的技巧描述、以及内容/背景说明（即"为什么这么答"）。
# 绝不重复/改写题目原文，也绝不给出/重写答案本身。
EXPLAIN_PROMPT = (
    "你是一名资深语文教研助手。用户会给一道语文题目（可能含多个小题），该题的「答案」已在原文档给出、单独提供给你参考。\n"
    "请你只针对这道题写出一条「解析」，说明这个答案的来由：\n"
    "1) 第一行写「解析：」，随后写出对本道题（含各小题）的解析：讲清答题采分点有哪些、各占多少、如何拿分（答题思路/要点），并结合必要的内容背景。\n"
    "2) 解析要具体、紧扣题目与答案，不要空泛套话；没有把握的内容可略写，切勿编造。\n"
    "3) 绝对不要重复题目原文，也不要重写或给出「答案：」——答案请交给原文档。只输出「解析：」这一条（可多行），不要解释、不要加引号包裹。"
)


def load_minimax_key():
    if yaml is None:
        return None, False, "缺少依赖 pyyaml（pip install pyyaml）"
    p = DEFAULT_CONFIG
    if not p.exists():
        alt = Path.home() / ".minimax" / "config.yaml"
        p = alt if alt.exists() else p
    if not p.exists():
        return None, False, f"未找到配置文件 {p}"
    try:
        data = yaml.safe_load(open(p, encoding="utf-8"))
    except Exception as e:
        return None, False, f"读取配置失败：{e}"
    opts = (data or {}).get("provider", {}).get("minimax", {}).get("options", {})
    key = (opts.get("apiKey", "") or "").strip()
    if not key or key.lower().startswith("sk-xxx") or len(key) < 10:
        return None, False, "MiniMax apiKey 仍是占位符（sk-xxx），请填真实 key 到 ~/.minimax/config.yaml"
    return key, True, "ok"


MODELS = ["MiniMax-M3", "MiniMax-M2.7", "abab6.5s-chat"]

def call_minimax(api_key, text, system, model_id=DEFAULT_MODEL_ID):
    if httpx is None:
        raise RuntimeError("缺少依赖 httpx（pip install httpx）")
    url = MINIMAX_CHINA_BASE + "/v1/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": "Bearer " + api_key}
    body = {"model": model_id, "max_tokens": 8192, "temperature": 0.3,
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": text or ""}]}
    with httpx.Client(timeout=300.0) as c:
        r = c.post(url, headers=headers, json=body)
        if r.status_code != 200:
            raise RuntimeError(f"LLM 接口返回 HTTP {r.status_code}: {r.text[:400]}")
        data = r.json()
    choices = data.get("choices") or []
    if not choices:
        raise RuntimeError("大模型未返回任何内容（choices 为空），可能是该段内容过长或接口波动，请重试，或拆成更小的讲次再分析。")
    content = (choices[0].get("message") or {}).get("content", "") or ""
    if not content.strip():
        raise RuntimeError("大模型返回了空内容，可能是接口波动，请重试。")
    return re.sub(r"<\s*think\b[^>]*>.*?<\s*/\s*think\s*>", "", content, flags=re.S | re.I).strip()


CONTENT_TYPES = {
    ".html": "text/html; charset=utf-8", ".js": "application/javascript; charset=utf-8",
    ".css": "text/css; charset=utf-8", ".json": "application/json; charset=utf-8",
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif",
    ".svg": "image/svg+xml", ".ico": "image/x-icon", ".woff2": "font/woff2",
    ".doc": "application/msword", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class Handler(BaseHTTPRequestHandler):
    server_version = "PyUnified/1.0"

    def _json(self, code, payload):
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type,Authorization,x-api-key")
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type,Authorization,x-api-key")
        self.end_headers()

    # 文件上传识别：读取 PDF/Word 文本层并返回提取文本（不做 OCR；扫描件/图片型 PDF 需先转文字版）
    def _handle_upload(self):
        import time
        from urllib.parse import parse_qs, urlparse
        q = parse_qs(urlparse(self.path).query)
        name = (q.get("name") or [""])[0]
        ext = (Path(name or "").suffix or "").lower()
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length) if length else b""
        if not raw:
            return self._json(400, {"error": "未收到文件内容"})
        try:
            from io import BytesIO
            if ext == ".pdf":
                from pypdf import PdfReader
                pr = PdfReader(BytesIO(raw))
                pages = []
                for pg in (getattr(pr, "pages", None) or []):
                    tx = ""
                    try:
                        tx = pg.extract_text() or ""
                    except Exception:
                        tx = ""
                    pages.append(tx)
                text = "\n\n".join([p for p in pages if str(p).strip()]) or "".join(pages)
                note = "PDF（文本层）" if pages and any(str(p).strip() for p in pages) else "PDF（未能提取到文本，可能是扫描/图片型 PDF，请另存为文字版后再传）"
            elif ext in (".docx", ".doc"):
                import docx
                d = docx.Document(BytesIO(raw))
                parts = [(p.text or "") for p in d.paragraphs]
                for tbl in list(d.tables or []):
                    for row in list(tbl.rows or []):
                        parts.append(" | ".join((c.text or "") for c in list(row.cells)))
                text = "\n".join(parts)
                note = "Word（.docx/.doc 由 python-docx 读取；旧二进制 .doc 建议另存为 .docx）"
            else:
                return self._json(400, {"error": "仅支持 PDF 或 Word(.docx)；旧版 .doc 请先另存为 .docx。未识别类型：" + (ext or name)})
        except Exception as e:
            msg = (str(e) or "").splitlines()[0] if str(e) else "解析失败"
            return self._json(500, {"error": "文件解析失败：" + msg})
        if not str(text).strip():
            return self._json(200, {"result": "", "note": note + "；未能提取到有效文字。"})
        return self._json(200, {"result": str(text), "note": note})

    def _load_uploaded_media(self):
        media_file = ROOT / "uploaded-media.js"
        images = []
        if media_file.exists():
          try:
            txt = media_file.read_text(encoding="utf-8")
            m = re.search(r"window\.uploadedMedia\s*=\s*(\{[\s\S]*?\})\s*;?\s*$", txt, re.S)
            if m:
              data = json.loads(m.group(1))
              images = data.get("images") or []
          except Exception:
            images = []
        return media_file, images

    def _save_uploaded_media(self, images):
        media_file = ROOT / "uploaded-media.js"
        payload = {"images": images}
        media_file.write_text("window.uploadedMedia = " + json.dumps(payload, ensure_ascii=False, indent=2) + ";\n", encoding="utf-8")

    def _handle_upload_image(self):
        import time
        from urllib.parse import parse_qs, urlparse
        q = parse_qs(urlparse(self.path).query)
        name = (q.get("name") or [""])[0]
        ext = (Path(name or "").suffix or "").lower()
        image_exts = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}
        if ext not in image_exts:
            return self._json(400, {"error": "图片上传仅支持图片格式（png/jpg/jpeg/gif/webp/bmp/svg）"})
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length) if length else b""
        if not raw:
            return self._json(400, {"error": "未收到文件内容"})
        upload_dir = ROOT / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        safe_base = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(name or "image").stem).strip("._-") or "image"
        fname = f"{safe_base}-{int(time.time() * 1000)}{ext}"
        (upload_dir / fname).write_bytes(raw)
        media_file, images = self._load_uploaded_media()
        record = {
            "id": "img-" + time.strftime("%Y%m%d%H%M%S"),
            "name": fname,
            "sourceName": name or fname,
            "url": f"/uploads/{fname}",
            "createdAt": time.strftime("%Y-%m-%d %H:%M"),
        }
        images.append(record)
        try:
            self._save_uploaded_media(images)
        except Exception as e:
            return self._json(500, {"error": "图片已上传，但记录保存失败：" + (str(e).splitlines()[0] if str(e) else "未知错误"), "url": record["url"]})
        return self._json(200, {"ok": True, "url": record["url"], "name": fname, "record": record, "count": len(images), "note": "图片已上传并保存记录"})

    # 保存清洗结果到"清洗知识库"数据文件（clean-knowledge-bank.js），并入知识库源
    def _handle_save_clean_kb(self):
        import time
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length) if length else b""
        try:
            payload = json.loads(raw.decode("utf-8") or "{}")
        except Exception:
            payload = {}
        title = str(payload.get("title") or "").strip()
        content = str(payload.get("content") or "").strip()
        if not content:
            return self._json(400, {"error": "知识库内容为空"})
        if not title:
            title = "清洗知识库项 " + time.strftime("%m-%d %H:%M")
        entry = {
            "id": "clean-" + time.strftime("%Y%m%d%H%M%S"),
            "title": title,
            "kind": "clean",
            "content": content,
            "createdAt": time.strftime("%Y-%m-%d %H:%M"),
        }
        sources = []
        kb_file = ROOT / "clean-knowledge-bank.js"
        if kb_file.exists():
            try:
                txt = kb_file.read_text(encoding="utf-8")
                m = re.search(r"window\.cleanKnowledgeBank\s*=\s*(\{[\s\S]*?\})\s*;?\s*$", txt, re.S)
                if m:
                    data = json.loads(m.group(1))
                    sources = data.get("sources") or []
            except Exception:
                sources = []
        sources.append(entry)
        out = "window.cleanKnowledgeBank = " + json.dumps({"name": "清洗知识库", "sources": sources}, ensure_ascii=False, indent=2) + ";\n"
        try:
            kb_file.write_text(out, encoding="utf-8")
        except Exception as e:
            return self._json(500, {"error": "保存失败：" + str(e).splitlines()[0] if str(e) else "保存失败"})
        return self._json(200, {"ok": True, "entry": entry, "count": len(sources)})

    # ===== 题目库 CRUD：JSON 文件存储 (question-bank.js -> window.questionBank = {questions:[...]}) =====
    def _load_questions(self):
        qf = ROOT / "question-bank.js"
        if qf.exists():
            try:
                txt = qf.read_text(encoding="utf-8")
                m = re.search(r"window\.questionBank\s*=\s*(\{[\s\S]*?\})\s*;?\s*$", txt, re.S)
                if m:
                    d = json.loads(m.group(1))
                    return d.get("questions") or []
            except Exception:
                pass
        return []

    def _save_questions(self, qs):
        qf = ROOT / "question-bank.js"
        out = "window.questionBank = " + json.dumps({"name": "题库", "questions": qs}, ensure_ascii=False, indent=2) + ";\n"
        qf.write_text(out, encoding="utf-8")

    def _handle_questions_GET(self):
        qs = self._load_questions()
        return self._json(200, {"ok": True, "questions": qs, "count": len(qs)})

    def _handle_questions_POST(self):
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length) if length else b""
        try:
            payload = json.loads(raw.decode("utf-8") or "{}")
        except Exception:
            return self._json(400, {"error": "JSON 无效"})
        action = str(payload.get("action") or "save")
        qs = self._load_questions()
        if action == "delete":
            qid = str(payload.get("id") or "")
            qs = [q for q in qs if str(q.get("id") or "") != qid]
            try:
                self._save_questions(qs)
                return self._json(200, {"ok": True, "count": len(qs)})
            except Exception as e:
                return self._json(500, {"error": "保存失败：" + str(e).splitlines()[0] if str(e) else "保存失败"})
        # save：单题(question) 或 批量(questions)，按 id upsert
        to_save = payload.get("questions") if isinstance(payload.get("questions"), list) else (payload.get("question") if isinstance(payload.get("question"), dict) else None)
        if to_save is None:
            return self._json(400, {"error": "缺少 question / questions 字段"})
        items = to_save if isinstance(to_save, list) else [to_save]
        saved = 0
        for q in items:
            if not isinstance(q, dict) or not q.get("id"):
                continue
            idx = next((i for i, x in enumerate(qs) if str(x.get("id")) == str(q.get("id"))), None)
            if idx is None:
                qs.append(q)
            else:
                qs[idx] = q
            saved += 1
        try:
            self._save_questions(qs)
            return self._json(200, {"ok": True, "count": len(qs), "saved": saved})
        except Exception as e:
            return self._json(500, {"error": "保存失败：" + str(e).splitlines()[0] if str(e) else "保存失败"})

    # ===== 试卷结构 CRUD：JSON 文件存储 (exam-structure-store.js -> window.examStructureStore = {papers:{...}}) =====
    def _load_exam_structures(self):
        sf = ROOT / "exam-structure-store.js"
        if sf.exists():
            try:
                txt = sf.read_text(encoding="utf-8")
                m = re.search(r"window\.examStructureStore\s*=\s*(\{[\s\S]*?\})\s*;?\s*$", txt, re.S)
                if m:
                    d = json.loads(m.group(1))
                    papers = d.get("papers")
                    return papers if isinstance(papers, dict) else {}
            except Exception:
                pass
        return {}

    def _save_exam_structures(self, papers):
        sf = ROOT / "exam-structure-store.js"
        out = "window.examStructureStore = " + json.dumps({"name": "试卷结构库", "papers": papers}, ensure_ascii=False, indent=2) + ";\n"
        sf.write_text(out, encoding="utf-8")

    def _handle_exam_structure_GET(self):
        from urllib.parse import parse_qs, urlparse
        papers = self._load_exam_structures()
        q = parse_qs(urlparse(self.path).query)
        paper_id = str((q.get("paperId") or [""])[0] or "")
        if paper_id:
            return self._json(200, {"ok": True, "paper": papers.get(paper_id), "paperId": paper_id})
        return self._json(200, {"ok": True, "papers": papers, "count": len(papers)})

    def _handle_exam_structure_POST(self):
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length) if length else b""
        try:
            payload = json.loads(raw.decode("utf-8") or "{}")
        except Exception:
            return self._json(400, {"error": "JSON 无效"})
        action = str(payload.get("action") or "save")
        paper_id = str(payload.get("paperId") or "").strip()
        if not paper_id:
            return self._json(400, {"error": "缺少 paperId"})
        papers = self._load_exam_structures()
        if action == "delete":
            papers.pop(paper_id, None)
            try:
                self._save_exam_structures(papers)
                return self._json(200, {"ok": True, "count": len(papers)})
            except Exception as e:
                return self._json(500, {"error": "保存失败：" + str(e).splitlines()[0] if str(e) else "保存失败"})
        paper = payload.get("paper")
        if not isinstance(paper, dict):
            return self._json(400, {"error": "缺少 paper 结构数据"})
        papers[paper_id] = paper
        try:
            self._save_exam_structures(papers)
            return self._json(200, {"ok": True, "paperId": paper_id, "count": len(papers)})
        except Exception as e:
            return self._json(500, {"error": "保存失败：" + str(e).splitlines()[0] if str(e) else "保存失败"})

    def do_GET(self):
        path = self.path.split("?")[0]
        if path == "/health":
            _k, ok, why = load_minimax_key()
            self._json(200, {"ok": True, "configured": ok, "note": why})
            return
        if path == "/api/questions":
            return self._handle_questions_GET()
        if path == "/api/exam-structure":
            return self._handle_exam_structure_GET()
        rel = (INDEX if path in ("/", "") else path.lstrip("/"))
        target = (ROOT / rel).resolve()
        if not str(target).startswith(str(ROOT)) or not target.is_file():
            self._json(404, {"error": "not found"})
            return
        data = target.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", CONTENT_TYPES.get(target.suffix.lower(), "application/octet-stream"))
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Access-Control-Allow-Origin", "*")
        # HTML/JS 禁用缓存，确保浏览器始终拿到最新代码（避免旧版仍指向已停的 8765）
        if target.suffix.lower() in (".html", ".js"):
            self.send_header("Cache-Control", "no-store, no-cache, must-revalidate, max-age=0")
            self.send_header("Pragma", "no-cache")
        self.end_headers()
        self.wfile.write(data)

    def do_POST(self):
        path0 = self.path.split("?")[0]
        if path0 == "/api/upload-image":
            return self._handle_upload_image()
        if path0 == "/api/upload":
            return self._handle_upload()
        if path0 == "/api/save-clean-kb":
            return self._handle_save_clean_kb()
        if path0 == "/api/questions":
            return self._handle_questions_POST()
        if path0 == "/api/exam-structure":
            return self._handle_exam_structure_POST()
        if path0 != "/api/clean":
            return self._json(404, {"error": "not found"})
        key, ok, why = load_minimax_key()
        if not ok:
            return self._json(503, {"error": why})
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length) if length else b""
        try:
            payload = json.loads(raw.decode("utf-8") or "{}")
        except Exception:
            payload = {}
        text = payload.get("text", "")
        if not text or not str(text).strip():
            return self._json(400, {"error": "请粘贴要清洗矫正的内容"})
        struct_mode = bool(payload.get("struct"))
        answers_mode = bool(payload.get("answers"))
        format_mode = bool(payload.get("format"))
        explain_mode = bool(payload.get("explain"))
        # 模式选择：format(仅规范化排版) > explain(仅补解析，答案已有) > answers(补单题答案/解析) > struct(重建题目+答案+解析) > 纯格式矫正
        if format_mode:
            sys_prompt = FORMAT_PROMPT
        elif explain_mode:
            sys_prompt = EXPLAIN_PROMPT
        elif answers_mode:
            sys_prompt = ANSWERS_PROMPT
        elif struct_mode:
            sys_prompt = STRUCT_PROMPT
        else:
            sys_prompt = SYSTEM_PROMPT
        # 多模型 fallback：M3 常只输出 think 而无最终答案，依次用 M3→M2.7→abab6.5s，取第一个非空内容
        result = ""
        used_model = DEFAULT_MODEL_ID
        last_err = ""
        for m in MODELS:
            try:
                r2 = call_minimax(key, text, sys_prompt, model_id=m)
            except Exception as e:
                last_err = str(e)
                continue
            if r2 and str(r2).strip():
                result = r2; used_model = m; break
        if not result or not str(result).strip():
            reason = last_err or "大模型对多个模型都未返回有效内容，请重试。"
            return self._json(502, {"error": reason})
        return self._json(200, {"result": result, "model": used_model, "struct": struct_mode, "answers": answers_mode, "format": format_mode, "explain": explain_mode})


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--port", type=int, default=8000)
    ap.add_argument("--host", default="0.0.0.0")
    args = ap.parse_args()
    srv = ThreadingHTTPServer((args.host, args.port), Handler)
    _k, ok, why = load_minimax_key()
    print(f"[server] 监听 http://{args.host}:{args.port}（静态页面 + /api/clean）")
    print(f"[server] MiniMax: {'已就绪' if ok else '未配置：' + why}")
    try:
        srv.serve_forever()
    except KeyboardInterrupt:
        print("\n[server] 已停止")


if __name__ == "__main__":
    main()

